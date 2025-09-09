import pandas as pd
from docx import Document
import os
from datetime import datetime
from dotenv import load_dotenv
import tempfile
from onedrive_uploader import OneDriveUploader
import msal
import requests
import time
import shutil

# Importar funciones de ConvertToPDF
try:
    from ConvertToPDF import get_all_word_documents, convert_word_to_pdf, extract_nit_from_filename
except ImportError:
    print("[WARNING] No se pudo importar ConvertToPDF.py")
    get_all_word_documents = None
    convert_word_to_pdf = None
    extract_nit_from_filename = None

# Cargar variables de entorno
print(f"[DEBUG] Directorio actual: {os.getcwd()}")
print(f"[DEBUG] Directorio del script: {os.path.dirname(__file__)}")

# Intentar cargar .env desde múltiples ubicaciones
env_paths = [
    os.path.join(os.path.dirname(__file__), '.env'),
    os.path.join(os.path.dirname(__file__), '..', '.env'),
    '.env'
]

env_loaded = False
for env_path in env_paths:
    if os.path.exists(env_path):
        load_dotenv(env_path)
        print(f"[DEBUG] Archivo .env cargado desde: {env_path}")
        env_loaded = True
        break

if not env_loaded:
    print(f"[WARNING] No se encontró archivo .env en ninguna ubicación: {env_paths}")
    load_dotenv()  # Cargar desde ubicación por defecto

# Variables globales
TEMPLATE_PATH = None  # Se descargará desde OneDrive

def clean_existing_documents_onedrive():
    """Limpiar documentos existentes en OneDrive antes de generar nuevos"""
    access_token = get_access_token()
    if not access_token:
        print("[WARNING] No se pudo obtener token para limpiar documentos")
        return
    
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    headers = {'Authorization': f'Bearer {access_token}'}
    
    # Verificar si existe la carpeta Documentos_Generados/Comunicados
    folder_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Generados/Comunicados"
    response = requests.get(folder_url, headers=headers)
    
    if response.status_code == 200:
        # Eliminar la carpeta completa
        delete_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Generados/Comunicados"
        delete_response = requests.delete(delete_url, headers=headers)
        if delete_response.status_code == 204:
            print("[INFO] Documentos existentes eliminados de OneDrive")
        else:
            print(f"[WARNING] Error eliminando documentos: {delete_response.text}")
    else:
        print("[INFO] No hay documentos existentes que eliminar")

# Variables globales
access_token = None
temp_excel_path = None

def get_access_token():
    """Obtener token de acceso usando MSAL"""
    try:
        client_id = os.environ.get('AZURE_CLIENT_ID')
        client_secret = os.environ.get('AZURE_CLIENT_SECRET')
        tenant_id = os.environ.get('AZURE_TENANT_ID')
        user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL')
        
        print(f"[DEBUG] Variables de entorno:")
        print(f"  CLIENT_ID: {'OK' if client_id else 'FALTA'}")
        print(f"  CLIENT_SECRET: {'OK' if client_secret else 'FALTA'}")
        print(f"  TENANT_ID: {'OK' if tenant_id else 'FALTA'}")
        print(f"  USER_EMAIL: {user_email}")
        
        if not all([client_id, client_secret, tenant_id]):
            print("[ERROR] Faltan credenciales de Azure en .env")
            return None
            
        app = msal.ConfidentialClientApplication(
            client_id,
            authority=f"https://login.microsoftonline.com/{tenant_id}",
            client_credential=client_secret
        )
        
        result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
        
        if "access_token" in result:
            return result["access_token"]
        else:
            print(f" Error obteniendo token: {result.get('error_description')}")
            return None
            
    except Exception as e:
        print(f" Error: {str(e)}")
        return None

def download_template_from_onedrive():
    """Descargar plantilla Word desde OneDrive carpeta Documentos_Merge"""
    global access_token
    
    if not access_token:
        access_token = get_access_token()
    
    if not access_token:
        raise ValueError("No se pudo obtener token de acceso de Azure")
    
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    
    # Buscar archivos Word en la carpeta Documentos_Merge
    search_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Merge:/children"
    
    headers = {'Authorization': f'Bearer {access_token}'}
    response = requests.get(search_url, headers=headers)
    
    if response.status_code != 200:
        raise FileNotFoundError(f"No se pudo acceder a la carpeta Documentos_Merge: {response.text}")
    
    files = response.json().get('value', [])
    
    # Mostrar todos los archivos disponibles para depuración
    print(f"[DEBUG] Archivos encontrados en Documentos_Merge:")
    for f in files:
        print(f"  - {f['name']}")
    
    # Buscar plantilla específica RENOVACION_202X_CLIENTE
    template_files = [f for f in files if f['name'].endswith('.docx') and 
                     'renovacion' in f['name'].lower() and 'cliente' in f['name'].lower() and 
                     'propuesta' not in f['name'].lower()]
    
    if not template_files:
        docx_files = [f['name'] for f in files if f['name'].endswith('.docx')]
        raise FileNotFoundError(f"No se encontró plantilla .docx en Documentos_Merge. Archivos .docx disponibles: {docx_files}")
    
    # Usar el primer archivo de plantilla encontrado
    template_file = template_files[0]
    print(f"[DEBUG] Usando plantilla: {template_file['name']}")
    download_url = template_file['@microsoft.graph.downloadUrl']
    
    print(f"[INFO] Descargando plantilla desde OneDrive: {template_file['name']}")
    
    # Descargar archivo
    file_response = requests.get(download_url)
    if file_response.status_code != 200:
        raise FileNotFoundError("No se pudo descargar la plantilla")
    
    # Guardar en archivo temporal
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_file.write(file_response.content)
    temp_file.close()
    
    return temp_file.name

def download_excel_from_onedrive():
    """Descargar Excel desde OneDrive carpeta Documentos_Merge"""
    global access_token
    
    if not access_token:
        access_token = get_access_token()
    
    if not access_token:
        raise ValueError("No se pudo obtener token de acceso de Azure")
    
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    
    # Buscar archivos Excel en la carpeta Documentos_Merge
    search_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Merge:/children"
    
    headers = {'Authorization': f'Bearer {access_token}'}
    response = requests.get(search_url, headers=headers)
    
    if response.status_code != 200:
        raise FileNotFoundError(f"No se pudo acceder a la carpeta Documentos_Merge: {response.text}")
    
    files = response.json().get('value', [])
    excel_files = [f for f in files if f['name'].endswith(('.xlsx', '.xls'))]
    
    if not excel_files:
        raise FileNotFoundError("No se encontraron archivos Excel en Documentos_Merge")
    
    # Usar el primer archivo Excel encontrado
    excel_file = excel_files[0]
    download_url = excel_file['@microsoft.graph.downloadUrl']
    
    print(f"[INFO] Descargando Excel desde OneDrive: {excel_file['name']}")
    
    # Descargar archivo
    file_response = requests.get(download_url)
    if file_response.status_code != 200:
        raise FileNotFoundError("No se pudo descargar el archivo Excel")
    
    # Guardar en archivo temporal con nombre fijo en carpeta backend/Temp
    temp_dir = os.path.join(os.path.dirname(__file__), '..', 'Temp')
    os.makedirs(temp_dir, exist_ok=True)
    temp_excel_name = os.path.join(temp_dir, 'excel_clientes_temp.xlsx')
    with open(temp_excel_name, 'wb') as temp_file:
        temp_file.write(file_response.content)
    
    return temp_excel_name

def get_excel_data():
    global access_token, temp_excel_path
    
    # Usar OneDrive como única fuente
    print("[INFO] Descargando Excel desde OneDrive...")
    temp_excel_path = download_excel_from_onedrive()
    
    if not temp_excel_path:
        raise FileNotFoundError("No se pudo descargar el Excel desde OneDrive")
    
    return pd.read_excel(temp_excel_path, engine='openpyxl')

def get_onedrive_link(file_path, access_token):
    """Obtener link compartible de OneDrive para un archivo"""
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    headers = {'Authorization': f'Bearer {access_token}'}
    
    # Obtener información del archivo
    file_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/{file_path}"
    response = requests.get(file_url, headers=headers)
    
    if response.status_code == 200:
        file_info = response.json()
        return file_info.get('webUrl', '')
    return ''

def create_public_sharing_link(file_path, access_token):
    """Crear link público de compartición para un archivo PDF"""
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }
    
    # Crear link de compartición público
    share_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/{file_path}:/createLink"
    
    payload = {
        "type": "view",
        "scope": "anonymous"
    }
    
    response = requests.post(share_url, headers=headers, json=payload)
    
    if response.status_code == 201:
        share_info = response.json()
        return share_info.get('link', {}).get('webUrl', '')
    else:
        print(f"[WARNING] Error creando link público: {response.text}")
        return ''



def upload_excel_to_onedrive():
    """Subir Excel actualizado de vuelta a OneDrive"""
    global access_token, temp_excel_path
    
    if not access_token or not temp_excel_path:
        return
    
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    
    # Obtener nombre del archivo Excel original
    search_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Merge:/children"
    headers = {'Authorization': f'Bearer {access_token}'}
    response = requests.get(search_url, headers=headers)
    
    if response.status_code == 200:
        files = response.json().get('value', [])
        excel_files = [f for f in files if f['name'].endswith(('.xlsx', '.xls'))]
        
        if excel_files:
            excel_name = excel_files[0]['name']
            
            # Subir archivo actualizado
            with open(temp_excel_path, 'rb') as file_content:
                upload_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Merge/{excel_name}:/content"
                upload_response = requests.put(upload_url, headers=headers, data=file_content)
                
                if upload_response.status_code in [200, 201]:
                    print(f"[INFO] Excel actualizado en OneDrive: {excel_name}")
                else:
                    print(f"[WARNING] Error actualizando Excel: {upload_response.text}")

# Limpiar documentos existentes en OneDrive
print("[INFO] Limpiando documentos existentes en OneDrive...")
clean_existing_documents_onedrive()

# Cargar Excel
df = get_excel_data()

# Descargar plantilla desde OneDrive
TEMPLATE_PATH = download_template_from_onedrive()
print(f"[INFO] Plantilla descargada: {TEMPLATE_PATH}")

# Función para reemplazar etiquetas preservando formato
def reemplazar_etiquetas(doc, datos):
    # Obtener fecha actual
    fecha_actual = datetime.now()
    meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
             'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
    
    # Crear diccionario con datos de fecha
    datos_fecha = {
        'día': str(fecha_actual.day),
        'Mes': meses[fecha_actual.month - 1],
        'año': str(fecha_actual.year)
    }
    
    # Extraer solo el primer nombre si existe el campo "Nombre Representante"
    if 'Nombre Representante' in datos:
        primer_nombre = str(datos['Nombre Representante']).split()[0]
        datos['Nombre'] = primer_nombre
    
    # Combinar datos de fecha con datos del Excel
    todos_datos = {**datos_fecha, **datos}
    
    # Función auxiliar para reemplazar en párrafos
    def reemplazar_en_parrafos(parrafos):
        for p in parrafos:
            texto_completo = p.text
            for key, value in todos_datos.items():
                etiqueta_angular = f"<{key}>"
                etiqueta_guillemet = f"« {key}»"
                
                if etiqueta_angular in texto_completo or etiqueta_guillemet in texto_completo:
                    # Reemplazar en el texto completo del párrafo
                    nuevo_texto = texto_completo.replace(etiqueta_angular, str(value))
                    nuevo_texto = nuevo_texto.replace(etiqueta_guillemet, str(value))
                    
                    # Si hubo cambios, limpiar el párrafo y agregar el nuevo texto
                    if nuevo_texto != texto_completo:
                        p.clear()
                        p.add_run(nuevo_texto)
                        texto_completo = nuevo_texto
    
    # Reemplazar en párrafos principales
    reemplazar_en_parrafos(doc.paragraphs)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_en_parrafos(cell.paragraphs)

# Generar un documento por cada fila del Excel
for idx, row in df.iterrows():
    # Obtener NIT (primera columna)
    nit = str(row.iloc[0])  # Primera columna del Excel
    
    # El NIT se usará para crear carpetas en OneDrive
    
    # Cargar plantilla para cada cliente
    doc = Document(TEMPLATE_PATH)
    reemplazar_etiquetas(doc, row.to_dict())

    # Guardar documento
    nombre_archivo = f"Comunicado_{datetime.now().year}_{nit}.docx"
    
    # Siempre subir a OneDrive
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        doc.save(temp_file.name)

        # Obtener token y subir
        access_token = get_access_token()
        if access_token:
            uploader = OneDriveUploader(access_token, user_upn="mcanas@novacorp-plus.com")
            folder_path = f"Documentos_Generados/Comunicados/{nit}"
            uploader.create_folder(folder_path)
            uploader.upload_file(temp_file.name, folder_path, nombre_archivo)
            
            print(f"[INFO] Documento Word generado para NIT {nit}")

        # Eliminar archivo temporal
        try:
            os.unlink(temp_file.name)
        except Exception as e:
            print(f"[WARNING] No se pudo eliminar archivo temporal: {e}")

def convert_documents_to_pdf():
    """Convertir todos los documentos Word generados a PDF"""
    print(f"[DEBUG] Funciones PDF disponibles:")
    print(f"  get_all_word_documents: {get_all_word_documents is not None}")
    print(f"  convert_word_to_pdf: {convert_word_to_pdf is not None}")
    print(f"  extract_nit_from_filename: {extract_nit_from_filename is not None}")
    
    if not all([get_all_word_documents, convert_word_to_pdf, extract_nit_from_filename]):
        print("[WARNING] Funciones de conversión PDF no disponibles")
        return
    
    print("\n[INFO] Iniciando conversión de documentos a PDF...")
    
    access_token = get_access_token()
    if not access_token:
        print("[ERROR] No se pudo obtener token para conversión PDF")
        return
    
    # Procesar solo carpeta de Comunicados
    carpeta = "Documentos_Generados/Comunicados"
    word_files = get_all_word_documents(access_token, carpeta)
    
    if not word_files:
        print(f"[INFO] No se encontraron documentos Word en {carpeta}")
        return
    
    print(f"[INFO] Encontrados {len(word_files)} documentos para convertir")
    total_convertidos = 0
    
    for word_file in word_files:
        # Obtener carpeta del archivo
        file_folder = "/".join(word_file['path'].split("/")[:-1])
        
        pdf_link = convert_word_to_pdf(
            access_token, 
            word_file['id'], 
            word_file['name'], 
            file_folder
        )
        
        if pdf_link:
            total_convertidos += 1
            
            # Extraer NIT y actualizar Excel
            nit = extract_nit_from_filename(word_file['name'], "Comunicados")
            
            if nit and temp_excel_path:
                # Actualizar Excel con link del PDF
                nit_column = df.columns[0]
                mask = df[nit_column].astype(str) == str(nit)
                
                if mask.any():
                    if 'Link_Comunicado' not in df.columns:
                        df['Link_Comunicado'] = ''
                    
                    # Crear link público para el PDF
                    public_link = create_public_sharing_link(word_file['path'], access_token)
                    final_link = public_link if public_link else pdf_link
                    
                    df.loc[mask, 'Link_Comunicado'] = final_link
                    print(f"[INFO] Link PDF actualizado para NIT {nit}")
        
        time.sleep(0.5)  # Pausa para evitar límites de API
    
    print(f"[INFO] Conversión completada: {total_convertidos} PDFs generados")
    
    # Guardar Excel actualizado con links de PDFs
    if temp_excel_path:
        df.to_excel(temp_excel_path, index=False, engine='openpyxl')
        upload_excel_to_onedrive()
        print("[INFO] Excel actualizado con links de PDFs")

# Convertir documentos a PDF antes de limpiar temporales
convert_documents_to_pdf()

# Limpiar archivos temporales
if temp_excel_path:
    try:
        os.unlink(temp_excel_path)
        print(f"[INFO] Archivo temporal Excel eliminado: {temp_excel_path}")
    except Exception as e:
        print(f"[WARNING] No se pudo eliminar archivo temporal Excel: {e}")

if TEMPLATE_PATH:
    try:
        os.unlink(TEMPLATE_PATH)
        print(f"[INFO] Archivo temporal plantilla eliminado: {TEMPLATE_PATH}")
    except Exception as e:
        print(f"[WARNING] No se pudo eliminar archivo temporal plantilla: {e}")
