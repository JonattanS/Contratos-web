import pandas as pd
from docx import Document
import os
from datetime import datetime
from dotenv import load_dotenv
import tempfile
from onedrive_uploader import OneDriveUploader
import msal
import requests
import time
import shutil

# Importar funciones de ConvertToPDF
try:
    from ConvertToPDF import get_all_word_documents, convert_word_to_pdf, extract_nit_from_filename
except ImportError:
    print("[WARNING] No se pudo importar ConvertToPDF.py")
    get_all_word_documents = None
    convert_word_to_pdf = None
    extract_nit_from_filename = None

# Cargar variables de entorno
load_dotenv()

# Variables globales
TEMPLATE_PATH1 = None  # Se descargará desde OneDrive
TEMPLATE_PATH2 = None  # Se descargará desde OneDrive

def clean_existing_documents_onedrive():
    """Limpiar documentos existentes en OneDrive antes de generar nuevos"""
    access_token = get_access_token()
    if not access_token:
        print("[WARNING] No se pudo obtener token para limpiar documentos")
        return
    
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    headers = {'Authorization': f'Bearer {access_token}'}
    
    # Verificar si existe la carpeta Documentos_Generados/Renovaciones
    folder_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Generados/Renovaciones"
    response = requests.get(folder_url, headers=headers)
    
    if response.status_code == 200:
        # Eliminar la carpeta completa
        delete_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Generados/Renovaciones"
        delete_response = requests.delete(delete_url, headers=headers)
        if delete_response.status_code == 204:
            print("[INFO] Documentos existentes eliminados de OneDrive")
        else:
            print(f"[WARNING] Error eliminando documentos: {delete_response.text}")
    else:
        print("[INFO] No hay documentos existentes que eliminar")

# Variables globales
access_token = None
temp_excel_path = None

def get_access_token():
    """Obtener token de acceso usando MSAL"""
    try:
        client_id = os.environ.get('AZURE_CLIENT_ID')
        client_secret = os.environ.get('AZURE_CLIENT_SECRET')
        tenant_id = os.environ.get('AZURE_TENANT_ID')
        
        if not all([client_id, client_secret, tenant_id]):
            print(" Faltan credenciales de Azure en .env")
            return None
            
        app = msal.ConfidentialClientApplication(
            client_id,
            authority=f"https://login.microsoftonline.com/{tenant_id}",
            client_credential=client_secret
        )
        
        result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
        
        if "access_token" in result:
            return result["access_token"]
        else:
            print(f" Error obteniendo token: {result.get('error_description')}")
            return None
            
    except Exception as e:
        print(f" Error: {str(e)}")
        return None

def download_templates_from_onedrive():
    """Descargar plantillas Word desde OneDrive carpeta Documentos_Merge"""
    global access_token
    
    if not access_token:
        access_token = get_access_token()
    
    if not access_token:
        raise ValueError("No se pudo obtener token de acceso de Azure")
    
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    
    # Buscar archivos Word en la carpeta Documentos_Merge
    search_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Merge:/children"
    
    headers = {'Authorization': f'Bearer {access_token}'}
    response = requests.get(search_url, headers=headers)
    
    if response.status_code != 200:
        raise FileNotFoundError(f"No se pudo acceder a la carpeta Documentos_Merge: {response.text}")
    
    files = response.json().get('value', [])
    
    # Buscar plantillas específicas PROPUESTA FE
    template1_files = [f for f in files if f['name'].endswith('.docx') and 
                      'propuesta' in f['name'].lower() and 'cliente1' in f['name'].lower()]
    template2_files = [f for f in files if f['name'].endswith('.docx') and 
                      'propuesta' in f['name'].lower() and 'cliente2' in f['name'].lower()]
    
    if not template1_files:
        raise FileNotFoundError("No se encontró plantilla CLIENTE1 en Documentos_Merge")
    if not template2_files:
        raise FileNotFoundError("No se encontró plantilla CLIENTE2 en Documentos_Merge")
    
    # Descargar ambas plantillas
    templates = {}
    for i, (template_files, name) in enumerate([(template1_files, 'CLIENTE1'), (template2_files, 'CLIENTE2')], 1):
        template_file = template_files[0]
        download_url = template_file['@microsoft.graph.downloadUrl']
        
        print(f"[INFO] Descargando plantilla {name} desde OneDrive: {template_file['name']}")
        
        # Descargar archivo
        file_response = requests.get(download_url)
        if file_response.status_code != 200:
            raise FileNotFoundError(f"No se pudo descargar la plantilla {name}")
        
        # Guardar en archivo temporal
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.write(file_response.content)
        temp_file.close()
        
        templates[i] = temp_file.name
    
    return templates[1], templates[2]

def download_excel_from_onedrive():
    """Descargar Excel desde OneDrive carpeta Documentos_Merge"""
    global access_token
    
    if not access_token:
        access_token = get_access_token()
    
    if not access_token:
        raise ValueError("No se pudo obtener token de acceso de Azure")
    
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    
    # Buscar archivos Excel en la carpeta Documentos_Merge
    search_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Merge:/children"
    
    headers = {'Authorization': f'Bearer {access_token}'}
    response = requests.get(search_url, headers=headers)
    
    if response.status_code != 200:
        raise FileNotFoundError(f"No se pudo acceder a la carpeta Documentos_Merge: {response.text}")
    
    files = response.json().get('value', [])
    excel_files = [f for f in files if f['name'].endswith(('.xlsx', '.xls'))]
    
    if not excel_files:
        raise FileNotFoundError("No se encontraron archivos Excel en Documentos_Merge")
    
    # Usar el primer archivo Excel encontrado
    excel_file = excel_files[0]
    download_url = excel_file['@microsoft.graph.downloadUrl']
    
    print(f"[INFO] Descargando Excel desde OneDrive: {excel_file['name']}")
    
    # Descargar archivo
    file_response = requests.get(download_url)
    if file_response.status_code != 200:
        raise FileNotFoundError("No se pudo descargar el archivo Excel")
    
    # Guardar en archivo temporal con nombre fijo en carpeta backend/Temp
    temp_dir = os.path.join(os.path.dirname(__file__), '..', 'Temp')
    os.makedirs(temp_dir, exist_ok=True)
    temp_excel_name = os.path.join(temp_dir, 'excel_clientes_temp.xlsx')
    with open(temp_excel_name, 'wb') as temp_file:
        temp_file.write(file_response.content)
    
    return temp_excel_name

def get_excel_data():
    global access_token, temp_excel_path
    
    # Usar OneDrive como única fuente
    print("[INFO] Descargando Excel desde OneDrive...")
    temp_excel_path = download_excel_from_onedrive()
    
    if not temp_excel_path:
        raise FileNotFoundError("No se pudo descargar el Excel desde OneDrive")
    
    return pd.read_excel(temp_excel_path, engine='openpyxl')

def get_onedrive_link(file_path, access_token):
    """Obtener link compartible de OneDrive para un archivo"""
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    headers = {'Authorization': f'Bearer {access_token}'}
    
    # Obtener información del archivo
    file_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/{file_path}"
    response = requests.get(file_url, headers=headers)
    
    if response.status_code == 200:
        file_info = response.json()
        return file_info.get('webUrl', '')
    return ''

def create_public_sharing_link(file_path, access_token):
    """Crear link público de compartición para un archivo PDF"""
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    headers = {
        'Authorization': f'Bearer {access_token}',
        'Content-Type': 'application/json'
    }
    
    # Crear link de compartición público
    share_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/{file_path}:/createLink"
    
    payload = {
        "type": "view",
        "scope": "anonymous"
    }
    
    response = requests.post(share_url, headers=headers, json=payload)
    
    if response.status_code == 201:
        share_info = response.json()
        return share_info.get('link', {}).get('webUrl', '')
    else:
        print(f"[WARNING] Error creando link público: {response.text}")
        return ''



def upload_excel_to_onedrive():
    """Subir Excel actualizado de vuelta a OneDrive"""
    global access_token, temp_excel_path
    
    if not access_token or not temp_excel_path:
        return
    
    user_email = os.environ.get('EXTERNAL_ONEDRIVE_EMAIL', 'mcanas@novacorp-plus.com')
    
    # Obtener nombre del archivo Excel original
    search_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Merge:/children"
    headers = {'Authorization': f'Bearer {access_token}'}
    response = requests.get(search_url, headers=headers)
    
    if response.status_code == 200:
        files = response.json().get('value', [])
        excel_files = [f for f in files if f['name'].endswith(('.xlsx', '.xls'))]
        
        if excel_files:
            excel_name = excel_files[0]['name']
            
            # Subir archivo actualizado
            with open(temp_excel_path, 'rb') as file_content:
                upload_url = f"https://graph.microsoft.com/v1.0/users/{user_email}/drive/root:/Documentos_Merge/{excel_name}:/content"
                upload_response = requests.put(upload_url, headers=headers, data=file_content)
                
                if upload_response.status_code in [200, 201]:
                    print(f"[INFO] Excel actualizado en OneDrive: {excel_name}")
                else:
                    print(f"[WARNING] Error actualizando Excel: {upload_response.text}")

# Limpiar documentos existentes en OneDrive
print("[INFO] Limpiando documentos existentes en OneDrive...")
clean_existing_documents_onedrive()

# Cargar Excel
df = get_excel_data()

# Descargar plantillas desde OneDrive
TEMPLATE_PATH1, TEMPLATE_PATH2 = download_templates_from_onedrive()
print(f"[INFO] Plantillas descargadas: {TEMPLATE_PATH1}, {TEMPLATE_PATH2}")

# Función para reemplazar etiquetas en párrafos y tablas
def reemplazar_etiquetas(doc, datos):
    fecha_actual = datetime.now()
    meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
             'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
    
    datos_fecha = {
        'día': str(fecha_actual.day),
        'Mes': meses[fecha_actual.month - 1],
        'AÑO': str(fecha_actual.year),  # AÑO en mayúscula porque en la plantilla lo está así
        'MES': meses[fecha_actual.month - 1].upper(),  # para el <MES> del pie de página si se escribe en mayúscula
        'año': str(fecha_actual.year),
        'mes': meses[fecha_actual.month - 1]
    }

    if 'Nombres y Apellidos' in datos:
        datos['Nombre'] = str(datos['Nombres y Apellidos']).split()[0]
    
    todos_datos = {**datos_fecha, **datos}

    def reemplazar_en_parrafos(parrafos):
        for p in parrafos:
            texto_completo = p.text
            for key, value in todos_datos.items():
                patrones = [
                    f"<{key}>",        # <Valor por Documento>
                    f"« {key}»",       # « Valor por Documento»
                    f"$<{key}>",       # $<Valor por Documento>
                    f"${{{key}}}",     # ${Valor por Documento}
                ]
                
                for etiqueta in patrones:
                    if etiqueta in texto_completo:
                        # Reemplazar en el texto completo del párrafo
                        nuevo_texto = texto_completo.replace(etiqueta, str(value))
                        
                        # Si hubo cambios, limpiar el párrafo y agregar el nuevo texto
                        if nuevo_texto != texto_completo:
                            p.clear()
                            p.add_run(nuevo_texto)
                            texto_completo = nuevo_texto

    def reemplazar_en_tablas(tablas):
        for table in tablas:
            for row in table.rows:
                for cell in row.cells:
                    reemplazar_en_parrafos(cell.paragraphs)

    reemplazar_en_parrafos(doc.paragraphs)
    reemplazar_en_tablas(doc.tables)

    #  Encabezado y pie de página
    for section in doc.sections:
        reemplazar_en_parrafos(section.footer.paragraphs)
        reemplazar_en_parrafos(section.header.paragraphs)


# Generar documentos
for idx, row in df.iterrows():
    nit = str(row.iloc[0])
    indicador = row.get('Indicador Tarifa', '')

    # Elegir plantilla según el valor numérico del indicador
    if pd.isna(indicador):
        print(f" Indicador vacío para NIT {nit}. Se usará plantilla 1 por defecto.")
        plantilla = TEMPLATE_PATH1
    elif str(indicador).strip() == "1":
        plantilla = TEMPLATE_PATH1
    elif str(indicador).strip() == "2":
        plantilla = TEMPLATE_PATH2
    else:
        print(f" Indicador no reconocido ('{indicador}') para NIT {nit}. Se usará plantilla 1 por defecto.")
        plantilla = TEMPLATE_PATH1

    # Cargar plantilla seleccionada
    doc = Document(plantilla)

    # El NIT se usará para crear carpetas en OneDrive

    # Agregar mapeo específico para RAZÓN SOCIAL
    datos_cliente = row.to_dict()
    if 'RAZÓN SOCIAL' in datos_cliente or 'Razón Social' in datos_cliente:
        # Usar el campo tal como está en el Excel
        pass
    elif 'Razon Social' in datos_cliente:
        datos_cliente['RAZÓN SOCIAL'] = datos_cliente['Razon Social']
    
    # Reemplazar etiquetas
    reemplazar_etiquetas(doc, datos_cliente)

    # Guardar documento
    nombre_archivo = f"Renovacion_{datetime.now().year}_{nit}.docx"
    
    # Siempre subir a OneDrive
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        doc.save(temp_file.name)

        # Obtener token y subir
        access_token = get_access_token()
        if access_token:
            uploader = OneDriveUploader(access_token, user_upn="mcanas@novacorp-plus.com")
            folder_path = f"Documentos_Generados/Renovaciones/{nit}"
            uploader.create_folder(folder_path)
            uploader.upload_file(temp_file.name, folder_path, nombre_archivo)
            
            # Actualizar Excel con link del Word
            word_path = f"{folder_path}/{nombre_archivo}"
            word_link = get_onedrive_link(word_path, access_token)
            
            if word_link:
                nit_column = df.columns[0]
                mask = df[nit_column].astype(str) == str(nit)
                
                if mask.any():
                    if 'Link_Renovacion' not in df.columns:
                        df['Link_Renovacion'] = ''
                    
                    df.loc[mask, 'Link_Renovacion'] = word_link
                    print(f"[INFO] Link Word actualizado para NIT {nit}")
            
            print(f"[INFO] Documento Word generado para NIT {nit}")

        # Eliminar archivo temporal
        try:
            os.unlink(temp_file.name)
        except Exception as e:
            print(f"[WARNING] No se pudo eliminar archivo temporal: {e}")

# Guardar Excel actualizado con links de Word
if temp_excel_path:
    df.to_excel(temp_excel_path, index=False, engine='openpyxl')
    upload_excel_to_onedrive()
    print("[INFO] Excel actualizado con links de Word")

# Limpiar archivos temporales
# Mantener Excel temporal para notificaciones
# if temp_excel_path:
#     try:
#         os.unlink(temp_excel_path)
#         print(f"[INFO] Archivo temporal Excel eliminado: {temp_excel_path}")
#     except Exception as e:
#         print(f"[WARNING] No se pudo eliminar archivo temporal Excel: {e}")

if TEMPLATE_PATH1:
    try:
        os.unlink(TEMPLATE_PATH1)
        print(f"[INFO] Archivo temporal plantilla 1 eliminado: {TEMPLATE_PATH1}")
    except Exception as e:
        print(f"[WARNING] No se pudo eliminar archivo temporal plantilla 1: {e}")

if TEMPLATE_PATH2:
    try:
        os.unlink(TEMPLATE_PATH2)
        print(f"[INFO] Archivo temporal plantilla 2 eliminado: {TEMPLATE_PATH2}")
    except Exception as e:
        print(f"[WARNING] No se pudo eliminar archivo temporal plantilla 2: {e}")
